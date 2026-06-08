import docx

for name in ["template 1.docx", "template 2.docx", "template 3.docx"]:
    doc = docx.Document(name)
    text = ""
    for p in doc.paragraphs:
        text += p.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    print(f"=== {name} ===")
    if "\u0e54.\u0e53" in text or "4.3" in text:
        print("Contains 4.3")
    if "\u0e54.\u0e54" in text or "4.4" in text:
        print("Contains 4.4")
