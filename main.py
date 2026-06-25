import streamlit as st
import auth
import pandas as pd
import gspread
from datetime import datetime, timedelta
import os
import io
import streamlit.components.v1 as components
import time
import requests
import base64
from docx import Document
import zipfile
import re

# PDF generation libraries
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import pink, lightgrey, blue
from pypdf import PdfReader, PdfWriter
import io
import os

try:
    pdfmetrics.registerFont(TTFont('THSarabunNew', 'THSarabunNew.ttf'))
    pdfmetrics.registerFont(TTFont('THSarabun', 'THSarabunNew.ttf')) # ลงทะเบียนชื่อ THSarabun เพิ่มเติม
except Exception as e:
    print(f"Font registration error: {e}")

def format_cid(cid):
    """Format Thai ID (13 digits) to: x xxxx xxxxx xx x"""
    cid = str(cid).strip().replace(' ', '').replace('-', '')
    if len(cid) == 13:
        return f"{cid[0]} {cid[1:5]} {cid[5:10]} {cid[10:12]} {cid[12]}"
    return cid

def num_to_thai_baht(num_str):
    try:
        num_str = str(num_str).replace(',', '').strip()
        if not num_str or num_str == '-':
            return '-'
        num = float(num_str)
    except ValueError:
        return '-'
    if num == 0:
        return 'ศูนย์บาทถ้วน'
    baht = int(num)
    satang = int(round((num - baht) * 100))
    def convert_group(val_str, is_first_group):
        thai_digits = ['', 'หนึ่ง', 'สอง', 'สาม', 'สี่', 'ห้า', 'หก', 'เจ็ด', 'แปด', 'เก้า']
        thai_positions = ['', 'สิบ', 'ร้อย', 'พัน', 'หมื่น', 'แสน']
        n = len(val_str)
        result = ''
        for i, char in enumerate(val_str):
            digit = int(char)
            pos = n - 1 - i
            if digit != 0:
                if pos == 1 and digit == 1:
                    result += 'สิบ'
                elif pos == 1 and digit == 2:
                    result += 'ยี่สิบ'
                elif pos == 0 and digit == 1:
                    if n > 1 or not is_first_group:
                        result += 'เอ็ด'
                    else:
                        result += 'หนึ่ง'
                else:
                    result += thai_digits[digit] + thai_positions[pos]
        return result
    baht_str = str(baht)
    millions_groups = []
    while baht_str:
        millions_groups.append(baht_str[-6:])
        baht_str = baht_str[:-6]
    baht_text = ''
    num_groups = len(millions_groups)
    for idx, group in enumerate(reversed(millions_groups)):
        is_first_group = (idx == 0)
        group_text = convert_group(group, is_first_group)
        if group_text:
            baht_text += group_text
            if idx < num_groups - 1:
                baht_text += 'ล้าน'
        else:
            if idx < num_groups - 1 and baht_text:
                baht_text += 'ล้าน'
    if baht_text:
        baht_text += 'บาท'
    satang_text = ''
    if satang > 0:
        satang_text = convert_group(str(satang), True) + 'สตางค์'
        if baht_text == '':
            baht_text = satang_text
        else:
            baht_text += satang_text
    else:
        if baht_text:
            baht_text += 'ถ้วน'
    return baht_text

def update_fee_text(index):
    fee_val = st.session_state.get(f"p_fee_{index}", "")
    st.session_state[f"p_fee_txt_{index}"] = num_to_thai_baht(fee_val)


def create_pdf_overlay(data):
    # แปลงตัวเลขทั้งหมดใน data ให้เป็นตัวเลขไทย
    def to_thai_numerals(v):
        if v is None: return ""
        return str(v).translate(str.maketrans("0123456789", "๐๑๒๓๔๕๖๗๘๙"))
    data = {k: to_thai_numerals(v) for k, v in data.items()}

    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=(595.27, 841.89)) # A4 size
    can.setFont('THSarabunNew', 10.0)
    
    base_h = 841.89
    
    # ดึงค่าเงื่อนไขมาเช็คเพื่อเลือก Template
    val_43 = str(data.get('p_43', '')).strip()
    val_44 = str(data.get('p_44', '')).strip()
    
    # 1. กำหนดค่าเริ่มต้น (Default) สำหรับทุก Template
    # ปรับพิกัด V.6 (ดีดขึ้น 1 บรรทัดจาก V.5)
    y_offset = 5.5
    name_x = 325
    cid_x = 165
    shop_x = 230
    receipt_y_diff = 396.0 # ปรับให้ตรงบรรทัดใบเสร็จ (Y=445)
    date_x = 310 # ขยับมาทางขวาหน่อยเพื่อให้กึ่งกลางช่อง
    special_x = 170 # ขยับมาทางซ้ายให้ตรงกับ 4.1, 4.2 (เดิม 225)
    type_x = 110 # ค่าเริ่มต้นบรรทัดที่ 6
    special_y_43 = base_h - 578
    special_y_44 = base_h - 600
    
    # 2. ตั้งค่าพิกัดเริ่มต้น (Base Coordinates) จากการจูน V.17
    # คุณสามารถก๊อปปี้พิกัดเหล่านี้ไปวางใน if/elif ด้านล่างเพื่อแยกจูนแต่ละไฟล์ได้ครับ
    y1, y2, y3, y4 = 638, 612, 592, 573
    y5, y6, y7, y8 = 549, 529, 512, 492
    y9, y10, y11, y12 = 468, 450, 239, 212
    
    # ค่าธรรมเนียม X (Default)
    fee_x = 240
    fee_text_x = 330

    if val_43 and not val_44:
        template_file = "template4.3.pdf"
        y1, y2, y3, y4 = 638, 612, 592, 573
        y5, y6, y7, y8 = 549, 529, 512, 492
        y9, y10, y11, y12 = 468, 450, 239, 212
        # คุณสามารถปรับแกน X เฉพาะของ 4.3 ได้ที่นี่ครับ:
        # name_x = 275  # ปรับตัวเลขนี้เพื่อขยับชื่อ
        # shop_x = 230  # ปรับตัวเลขนี้เพื่อขยับชื่อร้าน
    elif val_43 and val_44:
        template_file = "template4.2.pdf"
        y1, y2, y3, y4 = 638, 612, 592, 573
        y5, y6, y7, y8 = 549, 529, 512, 492
        y9, y10, y11, y12 = 468, 450, 220, 195
        # คุณสามารถปรับแกน X เฉพาะของ 4.2 ได้ที่นี่ครับ:
        # name_x = 275 
        # shop_x = 230
    else:
        template_file = "template.pdf"

    # 3. เริ่มการวาดข้อความโดยใช้พิกัดที่ถูกกำหนดด้านบน (Direct Grid Lock)

    # วาดบรรทัดที่ 1: เล่มที่ / เลขที่ / ปี
    can.drawString(90, y1, str(data.get('p_license_book', '')))
    
    rcpt_no = str(data.get('p_rcpt_no', '')).strip()
    license_no = rcpt_no if rcpt_no and rcpt_no != "-" else str(data.get('p_license_no', ''))
    can.drawString(155, y1, license_no)
    
    can.drawString(210, y1, str(data.get('p_license_year', '')))
    
    # วาดบรรทัดที่ 2: ชื่อ / สัญชาติ
    can.drawString(name_x, y2, str(data.get('p_name', '')))
    can.drawString(515, y2, str(data.get('p_nationality', '')))
    
    # วาดบรรทัดที่ 3: ที่อยู่เจ้าของ
    can.drawString(125, y3, str(data.get('p_addr', '')))
    can.drawString(205, y3, str(data.get('p_moo', '')))
    can.drawString(280, y3, str(data.get('p_tumbon', '')))
    can.drawString(400, y3, str(data.get('p_amphoe', ''))) # ขยับไป 420
    can.drawString(500, y3, str(data.get('p_province', ''))) # ขยับไป 535
    
    # วาดบรรทัดที่ 4: CID / โทรศัพท์
    can.drawString(cid_x, y4, format_cid(data.get('p_cid', '')))
    can.drawString(380, y4, str(data.get('p_phone', '')))
    
    # วาดบรรทัดที่ 5: ชื่อสถานประกอบการ
    can.drawString(shop_x, y5, str(data.get('p_shop', '')))
    
    # วาดบรรทัดที่ 6: ประเภทกิจการ
    can.drawString(type_x, y6, str(data.get('p_type', '')))
    
    # วาดบรรทัดที่ 7: ที่อยู่สถานประกอบการ
    can.drawString(120, y7, str(data.get('p_shop_addr', '')))
    can.drawString(205, y7, str(data.get('p_shop_moo', '')))
    can.drawString(280, y7, str(data.get('p_shop_tumbon', '')))
    can.drawString(400, y7, str(data.get('p_shop_amphoe', ''))) # กลับมา 400
    can.drawString(500, y7, str(data.get('p_shop_province', ''))) # กลับมา 500
    
    # วาดบรรทัดที่ 8: โทรศัพท์สถานประกอบการ
    can.drawString(150, y8, str(data.get('p_shop_phone', '')))
    
    # วาดบรรทัดที่ 9: ค่าธรรมเนียม
    can.drawString(fee_x, y9, str(data.get('p_fee', '')))
    can.drawString(fee_text_x, y9, str(data.get('p_fee_text', '')))
    
    # วาดบรรทัดที่ 10: ใบเสร็จรับเงิน
    rcpt_book = str(data.get('p_rcpt_book', '')).strip()
    
    can.drawString(170, y10, rcpt_book)
    can.drawString(365, y10, str(data.get('p_rcpt_date', '')))
    
    # วาดบรรทัดที่ 11: วันออกใบอนุญาต (ข้อ 5)
    can.drawString(date_x, y11, str(data.get('issue_day', '')))
    can.drawString(date_x + 60, y11, str(data.get('issue_month', '')))
    can.drawString(date_x + 155, y11, str(data.get('issue_year', '')))
    
    # วาดบรรทัดที่ 12: วันสิ้นอายุ (ข้อ 6)
    can.drawString(date_x, y12, str(data.get('expire_day', '')))
    can.drawString(date_x + 60, y12, str(data.get('expire_month', '')))
    can.drawString(date_x + 155, y12, str(data.get('expire_year', '')))
    
    def draw_safe_string(x, y, text, max_w):
        # คำนวณความยาวจริง (นับเฉพาะตัวที่กินพื้นที่)
        while text and pdfmetrics.stringWidth(text, 'THSarabunNew', 10.0) > max_w:
            text = text[:-1]
        can.drawString(x, y, text)

    if val_43:
        draw_safe_string(special_x, special_y_43, val_43, 595 - special_x - 20)
    if val_44:
        draw_safe_string(special_x, special_y_44, val_44, 595 - special_x - 20)
    
    can.save()
    packet.seek(0)
    
    new_pdf = PdfReader(packet)
    
    try:
        existing_pdf = PdfReader(open(template_file, "rb"))
    except Exception as e:
        existing_pdf = PdfReader(open("template.pdf", "rb"))
        
    output = PdfWriter()
    page = existing_pdf.pages[0]
    page.merge_page(new_pdf.pages[0])
    output.add_page(page)
    
    output_stream = io.BytesIO()
    output.write(output_stream)
    output_stream.seek(0)
    return output_stream

def replace_pattern_in_runs(runs, pattern, replacement, underline=None):
    run_texts = [r.text for r in runs]
    full_text = "".join(run_texts)
    import re
    matches = list(re.finditer(pattern, full_text))
    if not matches:
        return False
    for match in reversed(matches):
        start_char, end_char = match.span()
        char_map = []
        for r_idx, r_text in enumerate(run_texts):
            for c_idx in range(len(r_text)):
                char_map.append((r_idx, c_idx))
        if not char_map:
            continue
        start_run_idx, start_run_char = char_map[start_char]
        if end_char - 1 < len(char_map):
            end_run_idx, end_run_char = char_map[end_char - 1]
        else:
            end_run_idx, end_run_char = char_map[-1]
        if start_run_idx == end_run_idx:
            r = runs[start_run_idx]
            run_text_list = list(r.text)
            run_text_list[start_run_char : start_run_char + (end_char - start_char)] = list(replacement)
            r.text = "".join(run_text_list)
            if underline is not None:
                r.font.underline = underline
        else:
            r_start = runs[start_run_idx]
            start_text = r_start.text[:start_run_char] + replacement
            r_start.text = start_text
            if underline is not None:
                r_start.font.underline = underline
            for idx in range(start_run_idx + 1, end_run_idx):
                runs[idx].text = ""
            r_end = runs[end_run_idx]
            end_text = r_end.text[end_run_char + 1:]
            r_end.text = end_text
        run_texts = [r.text for r in runs]
    return True

def trim_dots_from_runs(runs, to_remove):
    import re
    for run in reversed(runs):
        if to_remove <= 0:
            break
        text = run.text
        if not text:
            continue
        matches = list(re.finditer(r'\.{3,}', text))
        if not matches:
            continue
        new_text = list(text)
        for match in reversed(matches):
            span_start, span_end = match.span()
            span_len = span_end - span_start
            rem = min(to_remove, span_len)
            del new_text[span_end - rem : span_end]
            to_remove -= rem
            if to_remove <= 0:
                break
        run.text = "".join(new_text)

def trim_trailing_tabs_if_long(paragraph):
    runs = paragraph.runs
    if not runs:
        return
    trailing_run_indices = []
    for i in range(len(runs) - 1, -1, -1):
        text = runs[i].text
        if text and any(c not in '\t ' for c in text):
            break
        trailing_run_indices.append(i)
    if not trailing_run_indices:
        return
    non_trailing_text = "".join([runs[i].text for i in range(len(runs)) if i not in trailing_run_indices])
    if len(non_trailing_text) >= 50:
        for idx in trailing_run_indices:
            runs[idx].text = ""

def create_docx_document(data):
    # แปลงตัวเลขทั้งหมดใน data ให้เป็นตัวเลขไทย
    def to_thai_numerals(v):
        if v is None: return ""
        return str(v).translate(str.maketrans("0123456789", "๐๑๒๓๔๕๖๗๘๙"))

    raw_data = {k: to_thai_numerals(v) for k, v in data.items()}

    # ดึงค่าเงื่อนไขมาเช็คเพื่อเลือก Template
    val_43 = str(data.get('p_43', '')).strip()
    val_44 = str(data.get('p_44', '')).strip()

    if val_43 and not val_44:
        template_file = "template 2.docx"
    elif val_43 and val_44:
        template_file = "template 3.docx"
    else:
        template_file = "template 1.docx"

    # เล่มใบเสร็จ (ไม่รวมเลขที่แล้ว)
    rcpt_book = raw_data.get("p_rcpt_book", "")
    rcpt_no = raw_data.get("p_rcpt_no", "")
    rcpt_combined = rcpt_book

    mapped_data = {
        "p_license_book": raw_data.get("p_license_book", ""),
        "p_license_no": rcpt_no if rcpt_no and rcpt_no != "-" else raw_data.get("p_license_no", ""),
        "p_license_year": raw_data.get("p_license_year", ""),
        "p_name": raw_data.get("p_name", ""),
        "p_nationality": raw_data.get("p_nationality", ""),
        "p_addr": raw_data.get("p_addr", ""),
        "p_moo": raw_data.get("p_moo", ""),
        "p_tumbon": raw_data.get("p_tumbon", ""),
        "p_amphoe": raw_data.get("p_amphoe", ""),
        "p_province": raw_data.get("p_province", ""),
        "p_cid": raw_data.get("p_cid", ""),
        "p_phone": raw_data.get("p_phone", ""),
        "p_shop": raw_data.get("p_shop", ""),
        "p_type": raw_data.get("p_type", ""),
        
        # ที่อยู่ร้าน
        "p_shopaddr": raw_data.get("p_shop_addr", raw_data.get("p_addr", "")),
        "p_shop_m": raw_data.get("p_shop_moo", raw_data.get("p_moo", "")),
        "p_shop_t": raw_data.get("p_shop_tumbon", raw_data.get("p_tumbon", "")),
        "p_shop_a": raw_data.get("p_shop_amphoe", raw_data.get("p_amphoe", "")),
        "p_shop_p": raw_data.get("p_shop_province", raw_data.get("p_province", "")),
        "p_shop_phone": raw_data.get("p_shop_phone", raw_data.get("p_phone", "")),
        
        # ค่าธรรมเนียม & ใบเสร็จ
        "p_fee": raw_data.get("p_fee", ""),
        "p_fee_text": raw_data.get("p_fee_text", ""),
        "p_rcpt_book": rcpt_combined,
        "p_rcpt_date": raw_data.get("p_rcpt_date", ""),
        
        # วันออกใบอนุญาต (issue)
        "d": raw_data.get("issue_day", ""),
        "m": raw_data.get("issue_month", ""),
        "y": raw_data.get("issue_year", ""),
        
        # วันหมดอายุ (expire)
        "d2": raw_data.get("expire_day", ""),
        "m2": raw_data.get("expire_month", ""),
        "y2": raw_data.get("expire_year", ""),
        
        # เงื่อนไข
        "p_43": val_43,
        "p_44": val_44,
    }

    doc = Document(template_file)

    def replace_placeholders(paragraphs):
        for paragraph in list(paragraphs):
            # ตรวจสอบและลบเงื่อนไขเมื่อไม่มีข้อมูล (ป้องกันสระ/หัวข้อ 4.3, 4.4 ค้างเปล่าๆ)
            if "๔.๓." in paragraph.text and not val_43:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                continue
            if "{{p_43}}" in paragraph.text and not val_43:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                continue
            if "๔.๔." in paragraph.text and not val_44:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                continue
            if "{{p_44}}" in paragraph.text and not val_44:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                continue

            orig_len = len(paragraph.text)
            has_changes = False

            for key, val in mapped_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    # ตรวจสอบความว่างเปล่า
                    is_empty = not val or str(val).strip() in ["", "-", "None"]
                    
                    if is_empty:
                        # ตรวจสอบว่ารอบๆ placeholder มีจุดไข่ปลาอยู่แล้วหรือไม่
                        has_dots = False
                        parts = paragraph.text.split(placeholder)
                        for i in range(len(parts) - 1):
                            before = parts[i]
                            after = parts[i+1]
                            before_dots = "." in before[-3:] if len(before) > 0 else False
                            after_dots = "." in after[:3] if len(after) > 0 else False
                            if before_dots or after_dots:
                                has_dots = True
                                break

                        if has_dots:
                            replacement_val = ""
                            pattern = re.escape(placeholder)
                        else:
                            # ในเทมเพลต Word (.docx) มีการใช้ปุ่ม Tab (\t) และตั้งค่า Tab Leader (เส้นประนำทาง) อยู่แล้ว
                            # การใส่จุดไข่ปลาจริงซ้ำซ้อนจะทำให้ตัวอักษรเลื่อนตกบรรทัด (Overflow)
                            # จึงใช้ค่าว่าง "" เพื่อให้ Tab Leader ของ Word วาดจุดไข่ปลาที่พอดีกับหน้ากระดาษโดยอัตโนมัติ
                            replacement_val = ""
                            pattern = re.escape(placeholder)
                    else:
                        val_str = str(val)
                        # เพิ่มช่องว่าง (Padding) รอบๆ ข้อมูลเพื่อให้แสดงเส้นประหรือระยะห่างสวยงามเหมือนในภาพ
                        if key in ["p_license_book", "p_license_no", "p_license_year"]:
                            replacement_val = f"   {val_str}   "
                        elif key in ["p_name", "p_shop", "p_type"]:
                            replacement_val = f"  {val_str}  "
                        elif key == "p_nationality":
                            replacement_val = f"  {val_str}  "
                        else:
                            replacement_val = val_str
                            
                        # หากมีข้อมูลจริง ให้กวาดล้างจุดไข่ปลาที่ขนาบข้าง placeholder ออกไป (เพื่อไม่ให้จุดโผล่ซ้ำ)
                        pattern = rf"\.*{re.escape(placeholder)}\.*"

                    if replace_pattern_in_runs(paragraph.runs, pattern, replacement_val):
                        has_changes = True

            if has_changes:
                new_len = len(paragraph.text)
                diff_len = new_len - orig_len
                if diff_len > 0:
                    trim_dots_from_runs(paragraph.runs, diff_len)

            trim_trailing_tabs_if_long(paragraph)

    replace_placeholders(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell.paragraphs)

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def create_app_docx_document(data):
    # แปลงตัวเลขทั้งหมดใน data ให้เป็นตัวเลขไทย
    def to_thai_numerals(v):
        if v is None: return ""
        if isinstance(v, bool): return v
        return str(v).translate(str.maketrans("0123456789", "๐๑๒๓๔๕๖๗๘๙"))
    
    raw_data = {k: to_thai_numerals(v) for k, v in data.items()}
    
    # แทนค่าช่อง checkbox
    mapped_data = {}
    for k, v in raw_data.items():
        if k.startswith("chk_"):
            mapped_data[k] = "☑" if v else "☐"
        else:
            mapped_data[k] = str(v) if v is not None else ""
            
    template_file = "แบบคำขอรับใบอนุญาต_ต่อใบอนุญาต add.docx"
    doc = Document(template_file)
    
    def replace_placeholders(paragraphs):
        for paragraph in list(paragraphs):
            orig_len = len(paragraph.text)
            has_changes = False
            
            for key, val in mapped_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    is_empty = not val or str(val).strip() in ["", "-", "None"]
                    if is_empty:
                        if key in ["p_road"]:
                            replacement_val = "    -    "
                        elif key in ["p_nationality"]:
                            replacement_val = "                    "  # เพิ่มเป็น 20 ช่องว่าง
                        elif key in ["p_agent", "p_evidence"]:
                            replacement_val = "                 "
                        elif key in ["p_soi"]:
                            replacement_val = "    -    "
                        elif key in ["p_name", "p_shop", "p_type", "p_fee_text", "p_rcpt_book"]:
                            replacement_val = "                              "
                        else:
                            replacement_val = "          "
                    else:
                        if key in ["p_road", "p_nationality"]:
                            replacement_val = f"  {str(val)}  "  
                        elif key in ["p_name", "p_shop", "p_type", "p_fee_text", "p_rcpt_book", "p_address", "p_moo", "p_soi"]:
                            replacement_val = f"      {str(val)}      "
                        else:
                            replacement_val = f" {str(val)} "
                    pattern = re.escape(placeholder)
                    
                    # 4 = WD_UNDERLINE.DOTTED
                    if replace_pattern_in_runs(paragraph.runs, pattern, replacement_val, underline=4):
                        has_changes = True
                        
            if has_changes:
                new_len = len(paragraph.text)
                diff_len = new_len - orig_len
                if diff_len > 0:
                    trim_dots_from_runs(paragraph.runs, diff_len)
                trim_trailing_tabs_if_long(paragraph)
                
    replace_placeholders(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell.paragraphs)
                
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

def create_app_pdf_overlay(data):
    # แปลงตัวเลขทั้งหมดใน data ให้เป็นตัวเลขไทย
    def to_thai_numerals(v):
        if v is None: return ""
        if isinstance(v, bool): return v
        return str(v).translate(str.maketrans("0123456789", "๐๑๒๓๔๕๖๗๘๙"))
    data = {k: to_thai_numerals(v) for k, v in data.items()}

    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=(595.27, 841.89)) # A4 size
    can.setFont('THSarabun', 10.0)
    
    # Use EXACT coordinates from the original document's words, minus a small offset to sit on the dotted line
    offset_y = -6.0
    
    can.drawString(260, 648.2 - offset_y, str(data.get('p_type', '')))
    can.drawString(360, 620.6 - offset_y, str(data.get('app_location', 'อบต.ดอยงาม')))
    
    can.drawString(355, 592.9 - offset_y, str(data.get('app_day', '')))
    can.drawString(415, 592.9 - offset_y, str(data.get('app_month', '')))
    can.drawString(505, 592.9 - offset_y, str(data.get('app_year', '')))
    
    can.drawString(150, 565.3 - offset_y, str(data.get('p_name', '')))
    can.drawString(355, 565.3 - offset_y, str(data.get('p_age', '')))
    can.drawString(430, 565.3 - offset_y, str(data.get('p_nationality', 'ไทย')))
    
    can.drawString(100, 537.6 - offset_y, str(data.get('p_agent', '')))
    
    can.drawString(205, 509.9 - offset_y, str(data.get('p_addr', '')))
    can.drawString(275, 509.9 - offset_y, str(data.get('p_moo', '')))
    
    p_soi_val = str(data.get('p_soi', '')).strip()
    p_road_val = str(data.get('p_road', '')).strip()
    can.drawString(350, 509.9 - offset_y, "-" if not p_soi_val or p_soi_val == "None" else p_soi_val)
    can.drawString(430, 509.9 - offset_y, "-" if not p_road_val or p_road_val == "None" else p_road_val)
    
    can.drawString(130, 482.3 - offset_y, str(data.get('p_subdistrict', 'ดอยงาม')))
    can.drawString(280, 482.3 - offset_y, str(data.get('p_district', 'พาน')))
    can.drawString(360, 482.3 - offset_y, str(data.get('p_province', 'เชียงราย')))
    
    can.drawString(150, 454.6 - offset_y, str(data.get('p_phone', '')))
    
    # Checkboxes Section 2
    if data.get('chk_1'): can.drawString(112, 399.3 - offset_y, "/")
    if data.get('chk_2'): can.drawString(112, 371.7 - offset_y, "/")
    if data.get('chk_3'): can.drawString(112, 344.1 - offset_y, "/")
    if data.get('chk_4'): can.drawString(112, 316.3 - offset_y, "/")
    if data.get('chk_5'): can.drawString(112, 288.7 - offset_y, "/")
    if data.get('chk_6'): can.drawString(112, 261.0 - offset_y, "/")
    if data.get('chk_7'): can.drawString(112, 233.4 - offset_y, "/")
    
    # Signature at bottom (moved left to prevent overlap)
    # can.drawString(215, 122.8 - offset_y, str(data.get('p_name', ''))) # Removed per user request to leave blank for manual signing
    
    can.save()
    packet.seek(0)
    
    new_pdf = PdfReader(packet)
    try:
        existing_pdf = PdfReader(open("app_template.pdf", "rb"))
    except FileNotFoundError:
        return None
        
    output = PdfWriter()
    page = existing_pdf.pages[0]
    page.merge_page(new_pdf.pages[0])
    output.add_page(page)
    
    output_stream = io.BytesIO()
    output.write(output_stream)
    output_stream.seek(0)
    return output_stream

# --- ส่วนที่ 1: การตั้งค่าการเชื่อมต่อและ API ---
SERVICE_ACCOUNT_FILE = 'credentials.json' 
SHEET_URL = "https://docs.google.com/spreadsheets/d/1t_xTI0RvXTUsUAJOvqgs1RFKsDK5f5qPCOPOvIqeMrw/edit"
GAS_URL = "https://script.google.com/macros/s/AKfycbxZ3Q1mk0hN69bmxctlOR95yKYC6hMP2BwaTXQn9WaX5edP09nUAKKv20N_kX-KzJg/exec"
@st.cache_resource
def get_gspread_client():
    try:
        scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        
        # สำหรับ Streamlit Cloud: ตรวจสอบจาก st.secrets
        if "gspread_credentials" in st.secrets:
            return gspread.service_account_from_dict(dict(st.secrets["gspread_credentials"]), scopes=scope)
        
        # ตรวจสอบว่ามีข้อมูลกุญแจอยู่ที่ top-level ของ secrets หรือไม่ (กรณีวางแบบไม่ครอบบล็อก)
        cert_dict = dict(st.secrets)
        if "private_key" in cert_dict:
            # ทำความสะอาด private_key (ลบช่องว่าง/จัดการ newline)
            pk = cert_dict["private_key"]
            if isinstance(pk, str):
                pk = pk.strip().replace("\\n", "\n")
            cert_dict["private_key"] = pk
            return gspread.service_account_from_dict(cert_dict, scopes=scope)
            
        # ถ้าไม่มีใน secrets (เช่น รันในเครื่องตัวเอง) ให้โหลดจากไฟล์
        if os.path.exists(SERVICE_ACCOUNT_FILE):
            return gspread.service_account(filename=SERVICE_ACCOUNT_FILE, scopes=scope)
        else:
            st.error(f"❌ ไม่พบข้อมูลการเชื่อมต่อทั้งใน st.secrets และไฟล์ {SERVICE_ACCOUNT_FILE}")
            return None
    except Exception as e:
        st.error(f"❌ โหลดไฟล์กุญแจ (JSON) ไม่สำเร็จ: {e}")
        return None

def update_gsheet(row_idx, data_dict, sheet_name=None):
    """Update a specific row in the Google Sheet based on row_idx (2-indexed)"""
    try:
        client = get_gspread_client()
        if not client: return False
        sh = client.open_by_url(SHEET_URL)
        
       
        if sheet_name:
            sheet = sh.worksheet(sheet_name)
        else:
            sheet = sh.get_worksheet(0)
            
     
        all_rows = sheet.get_all_values()
        if not all_rows: return False
        
        raw_headers = all_rows[0]
        headers = [h.strip() if h else "" for h in raw_headers]
        
 
        while headers and not headers[-1]:
            headers.pop()
            
        current_values = sheet.row_values(row_idx)
        
        # ปรับความยาว current_values ให้เท่ากับ headers
        if len(current_values) < len(headers):
            current_values.extend([""] * (len(headers) - len(current_values)))
        elif len(current_values) > len(headers):
            current_values = current_values[:len(headers)]
            
        update_values = []
        for i, h in enumerate(headers):
            if not h: # ถ้าหัวตารางว่าง ให้ข้ามหรือใช้ค่าเดิม
                update_values.append(current_values[i] if i < len(current_values) else "")
                continue
                
            if h in data_dict:
                update_values.append(data_dict[h])
            else:
                update_values.append(current_values[i] if i < len(current_values) else "")
        
        try:
            sheet.update(range_name=f"A{row_idx}", values=[update_values])
        except TypeError:
            sheet.update(f"A{row_idx}", [update_values])
        return True
    except Exception as e:
        st.error(f"❌ อัปเดต Google Sheet ไม่สำเร็จ: {e}")
        return False

def add_gsheet(data_dict, sheet_name=None):
    """Append a new row to the Google Sheet"""
    try:
        client = get_gspread_client()
        if not client: return False
        sh = client.open_by_url(SHEET_URL)
        
        # ใช้ sheet_name ที่ระบุ หรือถ้าไม่ระบุให้ใช้ sheet แรก
        if sheet_name:
            sheet = sh.worksheet(sheet_name)
        else:
            sheet = sh.get_worksheet(0)

        # หาหัวตาราง
        all_rows = sheet.get_all_values()
        if not all_rows: 
            st.error("❌ ไม่พบข้อมูลใดๆ ในชีตนี้")
            return False
            
        raw_headers = all_rows[0]
        headers = [h.strip() if h else "" for h in raw_headers]
        while headers and not headers[-1]:
            headers.pop()
            
        new_row = [data_dict.get(h, "") if h else "" for h in headers]
        
        # บันทึกข้อมูล
        sheet.append_row(new_row)
        return True
    except Exception as e:
        st.error(f"❌ เพิ่มข้อมูลลง Google Sheet ไม่สำเร็จ: {e}")
        return False

def upload_to_gdrive(file, folder_name, row_idx, sheet_name):
    """Upload file to Google Drive via Google Apps Script and update the sheet"""
    if not GAS_URL:
        st.error("⚠️ ยังไม่ได้ตั้งค่า GAS_URL กรุณาตั้งค่าในโค้ดก่อนใช้งาน")
        return False
        
    try:
        # Get Sheet ID from URL
        sheet_id = SHEET_URL.split("/d/")[1].split("/")[0]
        
        # Prepare file data
        file_bytes = file.read()
        file_b64 = base64.b64encode(file_bytes).decode('utf-8')
        
        payload = {
            "fileName": file.name,
            "fileData": file_b64,
            "mimeType": file.type,
            "folderName": folder_name,
            "sheetId": sheet_id,
            "sheetName": sheet_name,
            "rowIdx": row_idx # GAS is 1-indexed
        }
        
        with st.spinner("กำลังอัปโหลดไฟล์ไปยัง Google Drive..."):
            response = requests.post(GAS_URL, json=payload)
            result = response.json()
            
            if result.get("status") == "success":
                st.success(f"✅ อัปโหลดไฟล์ '{file.name}' สำเร็จ!")
                st.info(f"🔗 ลิงก์ไฟล์: {result.get('url')}")
                return True
            else:
                st.error(f"❌ อัปโหลดไม่สำเร็จ: {result.get('message')}")
                return False
    except Exception as e:
        st.error(f"❌ เกิดข้อผิดพลาดในการเชื่อมต่อ GAS: {e}")
        return False

def delete_gsheet_rows(row_indices, sheet_name=None):
    """Delete rows by indices (1-indexed list) in reverse order"""
    try:
        client = get_gspread_client()
        if not client: return False
        sh = client.open_by_url(SHEET_URL)
        if sheet_name:
            sheet = sh.worksheet(sheet_name)
        else:
            sheet = sh.get_worksheet(0)
            
        # ลบจากล่างขึ้นบนเพื่อป้องกัน index เลื่อน
        for idx in sorted(row_indices, reverse=True):
            sheet.delete_rows(idx)
        return True
    except Exception as e:
        st.error(f"❌ ลบข้อมูลไม่สำเร็จ: {e}")
        return False

def create_new_worksheet(new_name):
    """Create a new worksheet in the Google Sheet and copy headers from the first sheet"""
    try:
        client = get_gspread_client()
        if not client: 
            st.error("❌ ไม่สามารถเชื่อมต่อ Google Sheets ได้")
            return False
        sh = client.open_by_url(SHEET_URL)
        
        # ตรวจสอบว่าชื่อชีตซ้ำหรือไม่
        existing_names = [s.title for s in sh.worksheets()]
        if new_name in existing_names:
            st.error(f"⚠️ ชีตชื่อ '{new_name}' มีอยู่แล้วในระบบ")
            return False
            
        # ดึงหัวตารางจากแผ่นแรก
        first_sheet = sh.get_worksheet(0)
        headers = first_sheet.row_values(1) # อ่านแถวที่ 1
        
        # สร้างชีตใหม่ กำหนดแถว 1000 คอลัมน์เท่ากับหัวตารางหลัก
        num_cols = max(len(headers), 20)
        new_sheet = sh.add_worksheet(title=new_name, rows="1000", cols=str(num_cols))
        
        # เขียนหัวตารางลงไปในชีตใหม่
        new_sheet.append_row(headers)
        
        # เคลียร์ Cache ของ Streamlit เพื่อให้อัปเดตชื่อชีตทันที
        st.cache_data.clear()
        
        st.success(f"🎉 สร้างชีต '{new_name}' และคัดลอกหัวตารางสำเร็จเรียบร้อยแล้ว!")
        return True
    except Exception as e:
        st.error(f"❌ เกิดข้อผิดพลาดในการสร้างชีตใหม่: {e}")
        return False

def print_preview(title, content_html):
    html_template = f"""
    <html>
    <head>
        <link href="https://fonts.googleapis.com/css2?family=Sarabun&display=swap" rel="stylesheet">
        <style>
            body {{ font-family: 'Sarabun', sans-serif; padding: 20px; line-height: 1.5; }}
            .no-print {{ text-align: right; margin-bottom: 20px; }}
            .garuda {{ display: block; margin: 0 auto; width: 60px; }}
            .header {{ text-align: center; font-weight: bold; font-size: 18px; margin-top: 10px; }}
            .content {{ margin-top: 20px; font-size: 14px; }}
            @media print {{ .no-print {{ display: none; }} body {{ padding: 0; }} }}
            button {{ padding: 10px 20px; background-color: #1E3A8A; color: white; border: none; border-radius: 5px; cursor: pointer; }}
        </style>
    </head>
    <body>
        <div class="no-print">
            <button onclick="window.print()">🖨️ สั่งพิมพ์เอกสาร</button>
        </div>
        <img src="https://upload.wikimedia.org/wikipedia/commons/thumb/9/90/Garuda_Emb_of_Thailand_%28Black_and_White%29.svg/512px-Garuda_Emb_of_Thailand_%28Black_and_White%29.svg.png" class="garuda">
        <div class="header">{title}</div>
        <div class="content">{content_html}</div>
    </body>
    </html>
    """
    components.html(html_template, height=600, scrolling=True)

@st.cache_data(ttl=60)
def load_data(sheet_name=None):
    try:
        client = get_gspread_client()
        if not client: return pd.DataFrame(), {}, []
        sh = client.open_by_url(SHEET_URL)
        sheet_names = [s.title for s in sh.worksheets()]
        
        # เลือกชีตที่ต้องการโหลด ถ้าไม่ระบุให้เอาอันแรก
        if sheet_name and sheet_name in sheet_names:
            worksheet = sh.worksheet(sheet_name)
        else:
            worksheet = sh.get_worksheet(0)
            
        raw_data = worksheet.get_all_values()
        if not raw_data: return pd.DataFrame(), {}, sheet_names
        df = pd.DataFrame(raw_data[1:], columns=raw_data[0]).astype(str)
        df.columns = df.columns.str.strip()
        
        # จัดการชื่อคอลัมน์ที่ซ้ำกัน (เพื่อป้องกัน Error ใน Streamlit/Arrow)
        new_cols = []
        counts = {}
        for col in df.columns:
            if not col: col = "Unnamed"
            if col in counts:
                counts[col] += 1
                new_cols.append(f"{col}_{counts[col]}")
            else:
                counts[col] = 0
                new_cols.append(col)
        df.columns = new_cols
        
        # Helper to find column name among synonyms
        def find_col(synonyms, default):
            for s in synonyms:
                if s in df.columns: return s
            return default

        mapping = {
            'cid': find_col(['เลขประจำตัวประชาชน', 'เลขบัตรประชาชน'], 'เลขประจำตัวประชาชน'),
            'name': find_col(['ชื่อ - สกุล', 'ชื่อ-สกุล', 'ชื่อผู้ประกอบการ'], 'ชื่อ - สกุล'),
            'shop': find_col(['ชื่อสถานประกอบการ', 'ชื่อร้าน'], 'ชื่อสถานประกอบการ'),
            'expire': find_col(['วันหมดอายุใบอนุญาต', 'วันหมดอายุ'], 'วันหมดอายุใบอนุญาต'),
            'fee': find_col(['ค่าธรรมเนียม', 'จำนวนเงิน'], 'ค่าธรรมเนียม'),
            'type': find_col(['ประเภทการทำกิจการ', 'ประเภทกิจการ', 'ประเภท'], 'ประเภทกิจการ'),
            'address': find_col(['ที่อยู่เลขที่', 'สำนักงาน/บ้าน เลขที่', 'สำนักงาน/บ้านเลขที่', 'ที่อยู่/บ้านเลขที่', 'ตั้งอยู่เลขที่', 'บ้านเลขที่'], 'สำนักงาน/บ้านเลขที่'),
            'moo': find_col(['หมู่ที่', 'หมู่'], 'หมู่ที่'),
            'phone': find_col(['โทรศัพท์', 'เบอร์โทร', 'หมายเลขโทรศัพท์'], 'โทรศัพท์'),
            'rcpt_book': find_col(['ใบเสร็จรับเงินเล่มที่', 'เล่มที่'], 'ใบเสร็จรับเงินเล่มที่'),
            'rcpt_no': find_col(['เลขที่', 'เลขที่ใบเสร็จ'], 'เลขที่'),
            'rcpt_date': find_col(['ลงวันที่', 'วันที่รับเงิน'], 'ลงวันที่'),
            'attachment': find_col(['ไฟล์', 'ไฟล์แนบ', 'เอกสารแนบ', 'attachments'], 'ไฟล์'),
            'p_43': find_col(['เงื่อนไขพิเศษ4.3', 'เงื่อนไขพิเศษ 4.3', 'เงื่อนไข 4.3', 'เงื่อนไขเพิ่มเติม 4.3'], 'เงื่อนไขพิเศษ4.3'),
            'p_44': find_col(['เงื่อนไขพิเศษ4.4', 'เงื่อนไขพิเศษ 4.4', 'เงื่อนไข 4.4', 'เงื่อนไขเพิ่มเติม 4.4'], 'เงื่อนไขพิเศษ4.4')
        }
        if mapping['expire'] in df.columns:
            df[mapping['expire']] = pd.to_datetime(df[mapping['expire']], dayfirst=True, errors='coerce', format='mixed')
            df[mapping['expire']] = df[mapping['expire']].apply(
                lambda x: x.replace(year=x.year - 543) if pd.notnull(x) and x.year > 2500 else x
            )
        return df, mapping, sheet_names
    except Exception as e:
        import traceback
        st.error(f"❌ ไม่สามารถดึงข้อมูลได้ (รายละเอียด: {str(e)})")
        with st.expander("ดู Error แบบละเอียด (สำหรับ Debug)"):
            st.code(traceback.format_exc())
        return pd.DataFrame(), {}, []

@st.cache_data(ttl=60)
def load_sheet_data(sheet_name):
    try:
        client = get_gspread_client()
        if not client: return pd.DataFrame()
        sh = client.open_by_url(SHEET_URL)
        try:
            worksheet = sh.worksheet(sheet_name)
        except:
            return pd.DataFrame()
        raw_data = worksheet.get_all_values()
        if not raw_data: return pd.DataFrame()
        df = pd.DataFrame(raw_data[1:], columns=raw_data[0]).astype(str)
        df.columns = df.columns.str.strip()
        
        # จัดการชื่อคอลัมน์ที่ซ้ำกัน
        new_cols = []
        counts = {}
        for col in df.columns:
            if not col: col = "Unnamed"
            if col in counts:
                counts[col] += 1
                new_cols.append(f"{col}_{counts[col]}")
            else:
                counts[col] = 0
                new_cols.append(col)
        df.columns = new_cols
        return df
    except Exception as e:
        st.error(f"❌ ไม่สามารถดึงข้อมูลชีต {sheet_name} ได้: {str(e)}")
        return pd.DataFrame()

st.set_page_config(page_title="ระบบทะเบียนกิจการ อบต.ดอยงาม", layout="wide")

if not auth.check_login():
    st.stop()

# 1. ดึงรายชื่อชีตทั้งหมดมาก่อน (เพื่อใช้ใน Sidebar)
@st.cache_data(ttl=300)
def get_all_sheet_names():
    try:
        client = get_gspread_client()
        if not client: return ["รวมกิจการ"]
        sh = client.open_by_url(SHEET_URL)
        return [s.title for s in sh.worksheets()]
    except: return ["รวมกิจการ"]

sheet_names = get_all_sheet_names()

with st.sidebar:
    st.title("🏛️ อบต.ดอยงาม (Online)")
    
    # 2. ให้ผู้ใช้เลือกชีตเป้าหมาย
    target_sheet = st.selectbox("เลือกชีตเป้าหมาย (ออกตรวจ)", sheet_names)
    
    # ปุ่มสร้างชีตใหม่ (มีขอบเขต Expander สวยงาม)
    with st.sidebar.expander("➕ สร้างรอบออกตรวจใหม่"):
        with st.form("create_new_sheet_form"):
            new_sheet_name = st.text_input("ชื่อรอบการตรวจใหม่", placeholder="เช่น ออกตรวจ 15.6.68")
            submit_create = st.form_submit_button("สร้างรอบตรวจใหม่", use_container_width=True, type="primary")
            if submit_create:
                if not new_sheet_name.strip():
                    st.error("⚠️ กรุณาระบุชื่อชีตใหม่")
                else:
                    success = create_new_worksheet(new_sheet_name.strip())
                    if success:
                        st.rerun()
                        
    # 3. โหลดข้อมูลจากชีตที่เลือกมาใช้งาน
    df, cols, _ = load_data(target_sheet)
    
    app_category = st.selectbox("เลือกประเภทกิจการ", ["ทั้งหมด", "สถานประกอบกิจการ", "จำหน่าย/สะสมอาหาร", "ตลาด"])
    
    # --- กรองประเภทกิจการ (Global Filter) ---
    if app_category != "ทั้งหมด" and not df.empty:
        search_terms = []
        if app_category == "สถานประกอบกิจการ":
            search_terms = ["สถานประกอบ"]
        elif app_category == "จำหน่าย/สะสมอาหาร":
            search_terms = ["จำหน่าย", "สะสมอาหาร"]
        else:
            search_terms = [app_category]
            
        type_cols = []
        possible_cols = ['ประเภทกิจการ', 'ประเภทการทำกิจการ', 'ประเภท', cols.get('type')]
        for pc in possible_cols:
            if pc and pc in df.columns and pc not in type_cols:
                type_cols.append(pc)
        
        if type_cols:
            mask = pd.Series([False] * len(df), index=df.index)
            for c in type_cols:
                for term in search_terms:
                    mask |= df[c].str.contains(term, na=False)
            df = df[mask].copy()

    # ส่วนสำหรับ Debug
    with st.expander("🛠️ ตรวจสอบหัวตาราง (Debug)"):
        st.write(f"ชีตปัจจุบัน: {target_sheet}")
        st.caption("Version: V.35 (Max 100 visual chars)")
        if st.button("ล้างแคชและโหลดใหม่"):
            st.cache_data.clear()
            st.rerun()
    st.divider()
    menu = st.radio("เมนูหลัก", ["หน้าแรก (Dashboard)", "ค้นหา/จัดการข้อมูล"])
    if st.button("🔄 อัปเดตข้อมูลใหม่"):
        st.cache_data.clear()
        st.rerun()
    st.divider()
    if st.button("🚪 ออกจากระบบ", use_container_width=True):
        st.session_state.clear()
        st.rerun()

# --- หน้า Dashboard ---
if menu == "หน้าแรก (Dashboard)":
    st.header(f"📊 สรุปข้อมูล: {app_category}")
    if not df.empty:
        # ใช้ df ที่ถูกกรองมาแล้วจากระดับ Global
        f_df = df
            
        expired_df = pd.DataFrame()
        near_exp_df = pd.DataFrame()
        
        # กรองข้อมูลวันหมดอายุ (ถ้ามี)
        if cols['expire'] in f_df.columns:
            # แปลงวันที่ถ้ายังไม่ได้แปลง
            if not pd.api.types.is_datetime64_any_dtype(f_df[cols['expire']]):
                f_df[cols['expire']] = pd.to_datetime(f_df[cols['expire']], dayfirst=True, errors='coerce', format='mixed')
                
            today = datetime.now()
            near_date = today + timedelta(days=90)
            near_exp_df = f_df[(f_df[cols['expire']] >= today) & (f_df[cols['expire']] <= near_date)]
            expired_df = f_df[f_df[cols['expire']] < today]
        
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("ผู้ประกอบกิจการ", f"{f_df[cols['cid']].nunique() if cols['cid'] in f_df.columns else 0}")
        c2.metric("ใบอนุญาตทั้งหมด", f"{len(f_df)}")
        c3.metric("หมดอายุแล้ว", f"{len(expired_df)}", delta_color="inverse")
        c4.metric("ใกล้หมดอายุ (90 วัน)", f"{len(near_exp_df)}")
        
        st.divider()
        st.subheader(f"📋 รายชื่อผู้ประกอบการในชีต: {target_sheet}")
        st.info(f"พบข้อมูลทั้งหมด {len(f_df)} รายการ (กรองตามประเภท: {app_category})")
        
        # Display the dataframe
        st.dataframe(
            f_df, 
            use_container_width=True, 
            hide_index=True,
            column_config={
                "ลำดับที่": None,
                "ลำดับ": None
            }
        )
        
        # Add download button
        import io
        buffer = io.BytesIO()
        try:
            with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                safe_sheet_name = target_sheet[:31]
                export_df = f_df.drop(columns=[col for col in ['ลำดับ', 'ลำดับที่'] if col in f_df.columns])
                export_df.to_excel(writer, index=False, sheet_name=safe_sheet_name)
            
            st.download_button(
                label="📥 ดาวน์โหลดรายชื่อนี้เป็นไฟล์ Excel",
                data=buffer.getvalue(),
                file_name=f"รายชื่อ_{target_sheet}_{datetime.now().strftime('%d%m%Y')}.xlsx",
                mime="application/vnd.ms-excel",
                type="primary",
                key="dl_dash"
            )
        except Exception as e:
            pass

    # --- ส่วนดาวน์โหลดแบบตรวจกิจการ (ย้ายมาไว้ล่างสุดตามคำขอผู้ใช้) ---
    st.divider()
    with st.container(border=True):
        col_d1, col_d2 = st.columns([3, 1])
        with col_d1:
            st.markdown("""
            <div style="display: flex; align-items: center; gap: 10px;">
                <span style="font-size: 24px;">📋</span>
                <div style="line-height: 1.2;">
                    <h4 style="margin: 0; color: #1E3A8A; font-weight: bold;">ดาวน์โหลดแบบตรวจกิจการเปล่า (PDF)</h4>
                    <span style="font-size: 13px; color: #6B7280;">สำหรับนำไปใช้ในการออกตรวจหน้างานภายนอก</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
        with col_d2:
            inspection_file_path = "แบบตรวจกิจการ (ปรับปรุง).pdf"
            if os.path.exists(inspection_file_path):
                with open(inspection_file_path, "rb") as f:
                    st.download_button(
                        label="📥 ดาวน์โหลดแบบตรวจ (PDF)",
                        data=f,
                        file_name="แบบตรวจกิจการ_อบต_ดอยงาม.pdf",
                        mime="application/pdf",
                        key="inspection_btn_dashboard",
                        type="primary",
                        use_container_width=True
                    )
            else:
                st.error("❌ ไม่พบไฟล์แบบตรวจ")

# --- หน้าค้นหา/จัดการข้อมูล ---
elif menu == "ค้นหา/จัดการข้อมูล":
    st.header(f"🔍 จัดการข้อมูล (เป้าหมาย: {target_sheet})")
    
    # ส่วนแจ้งเตือนใกล้หมดอายุ (ย้ายมาไว้ที่นี่)
    with st.expander("🔔 ตรวจสอบรายชื่อใกล้หมดอายุ และที่หมดอายุแล้ว", expanded=False):
        f_df = df.copy()
        if cols['type'] in df.columns and app_category != "ทั้งหมด":
            f_df = df[df[cols['type']].str.contains(app_category, na=False)].copy()
            
        expired_df = pd.DataFrame()
        near_exp_df = pd.DataFrame()
        
        if cols['expire'] in f_df.columns:
            # แปลงวันที่ถ้ายังไม่ได้แปลง
            if not pd.api.types.is_datetime64_any_dtype(f_df[cols['expire']]):
                f_df[cols['expire']] = pd.to_datetime(f_df[cols['expire']], dayfirst=True, errors='coerce', format='mixed')
                
            today = datetime.now()
            near_date = today + timedelta(days=90)
            near_exp_df = f_df[(f_df[cols['expire']] >= today) & (f_df[cols['expire']] <= near_date)]
            expired_df = f_df[f_df[cols['expire']] < today]
        else:
            st.info("ℹ️ ชีตนี้ไม่มีคอลัมน์วันหมดอายุ ระบบจึงไม่สามารถแจ้งเตือนได้")
        
        if len(near_exp_df) > 0 or len(expired_df) > 0:
            if len(expired_df) > 0:
                st.error(f"⚠️ พบรายชื่อที่หมดอายุแล้วจำนวน {len(expired_df)} ราย")
            if len(near_exp_df) > 0:
                st.warning(f"🔔 พบรายชื่อใกล้หมดอายุ (90 วัน) จำนวน {len(near_exp_df)} ราย")
                
            import io
            buffer = io.BytesIO()
            try:
                # รวมทั้งสองกลุ่มเพื่อ Export
                combined_exp_df = pd.concat([expired_df, near_exp_df])
                
                with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                    potential = {
                        'name': [cols['name'], 'ชื่อ - สกุล'],
                        'cid': [cols['cid'], 'เลขประจำตัวประชาชน'],
                        'shop': [cols['shop'], 'ชื่อสถานประกอบการ'],
                        'address': [cols['address'], 'สำนักงาน/บ้าน เลขที่', 'สำนักงาน/บ้านเลขที่', 'บ้านเลขที่'],
                        'moo': [cols['moo'], 'หมู่ที่'],
                        'expire': [cols['expire'], 'วันหมดอายุใบอนุญาต']
                    }
                    final_cols = []
                    for k, v in potential.items():
                        for alias in v:
                            if alias in combined_exp_df.columns:
                                final_cols.append(alias); break
                    
                    export_df = combined_exp_df[final_cols].copy()
                    if cols['expire'] in export_df.columns:
                        export_df[cols['expire']] = export_df[cols['expire']].dt.strftime('%d/%m/%Y')
                    export_df.to_excel(writer, index=False, sheet_name='รายชื่อที่ต้องแจ้งเตือน')
                
                st.download_button("📥 ดาวน์โหลดไฟล์รายชื่อทั้งหมด (Excel)", buffer.getvalue(), 
                                 file_name=f"รายชื่อแจ้งเตือน_รวม_{datetime.now().strftime('%d%m%Y')}.xlsx",
                                 mime="application/vnd.ms-excel", type="primary")
                
                if len(expired_df) > 0:
                    st.subheader("🛑 รายชื่อที่หมดอายุแล้ว")
                    st.dataframe(
                        expired_df, 
                        use_container_width=True,
                        hide_index=True,
                        column_config={"ลำดับที่": None, "ลำดับ": None}
                    )
                if len(near_exp_df) > 0:
                    st.subheader("⏳ รายชื่อใกล้หมดอายุ")
                    st.dataframe(
                        near_exp_df, 
                        use_container_width=True,
                        hide_index=True,
                        column_config={"ลำดับที่": None, "ลำดับ": None}
                    )
            except Exception as e: st.error(f"สร้างไฟล์ไม่สำเร็จ: {e}")
        else: st.success("✅ ไม่มีรายชื่อใกล้หมดอายุ หรือหมดอายุแล้ว")

    st.divider()
    if st.button("➕ ลงทะเบียนรายใหม่", use_container_width=True, type="primary"):
        st.session_state['show_add_form'] = not st.session_state.get('show_add_form', False)

    # ฟอร์มลงทะเบียนรายใหม่ (ปรับปรุงให้ดึงตามคอลัมน์ในชีตจริง)
    if st.session_state.get('show_add_form', False):
        with st.form("add_new"):
            st.subheader("➕ ลงทะเบียนผู้ประกอบการใหม่")
            
            # เพิ่มตัวเลือกให้เลือกว่าจะลงออกตรวจวันไหน (Sheet ไหน)
            target_add_sheet = st.selectbox("เลือกวันที่ออกตรวจ (Sheet เป้าหมาย)", sheet_names, index=0)
            
            st.divider()
            c1, c2 = st.columns(2)
            new_entry = {}
            
            # ดึงหัวตารางจาก Sheet ที่เลือกเพื่อให้กรอกข้อมูลได้ถูกต้อง
            try:
                temp_client = get_gspread_client()
                temp_sh = temp_client.open_by_url(SHEET_URL)
                temp_ws = temp_sh.worksheet(target_add_sheet)
                fields = [h.strip() for h in temp_ws.row_values(1)]
            except:
                fields = df.columns.tolist()
            
            # คำนวณลำดับถัดไป (อิงตามข้อมูลเต็มของชีตเป้าหมาย ไม่โดน Filter)
            next_seq = "1"
            full_df = load_sheet_data(target_add_sheet)
            if not full_df.empty and 'ลำดับ' in full_df.columns:
                try:
                    seqs = pd.to_numeric(full_df['ลำดับ'], errors='coerce').dropna()
                    if not seqs.empty:
                        next_seq = str(int(seqs.max()) + 1)
                except: pass

            # กรองฟิลด์ที่ไม่ต้องการแสดงผล (Unnamed หรือไฟล์)
            visible_fields = [f for f in fields if f and not f.startswith('Unnamed') and f not in ['ไฟล์', 'ไฟล์แนบ']]
            
            for i, f in enumerate(visible_fields):
                t = c1 if i % 2 == 0 else c2
                if f == 'ลำดับ':
                    new_entry[f] = t.text_input(f, value=next_seq)
                elif 'วัน' in f:
                    new_entry[f] = t.date_input(f, value=datetime.now()).strftime('%d/%m/%Y')
                elif f == 'ตำบล' or f == 'ตำบล (ร้าน)':
                    new_entry[f] = t.text_input(f, value="ดอยงาม")
                elif f == 'อำเภอ' or f == 'อำเภอ (ร้าน)':
                    new_entry[f] = t.text_input(f, value="พาน")
                elif f == 'จังหวัด' or f == 'จังหวัด (ร้าน)':
                    new_entry[f] = t.text_input(f, value="เชียงราย")
                else:
                    new_entry[f] = t.text_input(f, value="")

            col_submit, col_cancel = st.columns(2)
            with col_submit:
                submitted = st.form_submit_button("💾 บันทึกข้อมูล", type="primary", use_container_width=True)
            with col_cancel:
                canceled = st.form_submit_button("❌ ยกเลิก", use_container_width=True)

            if submitted:
                with st.spinner(f"กำลังเพิ่มข้อมูลลงชีต {target_add_sheet}..."):
                    if add_gsheet(new_entry, sheet_name=target_add_sheet):
                        st.success(f"✅ เพิ่มข้อมูลลงชีต '{target_add_sheet}' สำเร็จ!")
                        time.sleep(1)
                        st.session_state['show_add_form'] = False
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.error("❌ บันทึกไม่สำเร็จ กรุณาลองใหม่อีกครั้ง")
            
            if canceled:
                st.session_state['show_add_form'] = False
                st.rerun()

    # ส่วนค้นหา
    col_search1, col_search2 = st.columns([8, 1])
    search = col_search1.text_input("🔍 ค้นหาชื่อ หรือ เลขบัตร...", placeholder="เช่น 3570500xxxxxx")
    
    with col_search2:
        st.markdown("<div style='margin-top: 27px;'></div>", unsafe_allow_html=True)
        search_btn = st.button("ค้นหา", use_container_width=True)
        
    if search or search_btn:
        if search:
            search_term = search.strip().replace('-', '').replace(' ', '')
        
        # 1. ลองค้นหาจากคอลัมน์หลักก่อน (ชื่อ และ เลขบัตร)
        mask = pd.Series(False, index=df.index)
        
        if cols['name'] in df.columns:
            mask |= df[cols['name']].str.replace(' ', '').str.contains(search_term, na=False)
        if cols['cid'] in df.columns:
            mask |= df[cols['cid']].str.replace('-', '').str.replace(' ', '').str.contains(search_term, na=False)
            
        results = df[mask]
        
        # 2. ถ้าไม่เจอ ให้ลองค้นหาจาก "ทุกคอลัมน์" (Global Search)
        if results.empty:
            results = df[df.apply(lambda row: row.astype(str).str.replace(' ', '').str.contains(search_term, na=False).any(), axis=1)]
            if not results.empty:
                st.info("💡 หมายเหตุ: พบข้อมูลจากการค้นหาทุกคอลัมน์ (หัวตารางอาจไม่ตรงตามมาตรฐาน)")
        
        if not results.empty:
            grouped = results.groupby(cols['name'], dropna=False)
            
            for person_name, group_data in grouped:
                first_row = group_data.iloc[0]
                u_name = first_row.get(cols['name'], 'ไม่ระบุชื่อ')
                u_cid = first_row.get(cols['cid'], '-')
                u_addr = first_row.get(cols['address'], '-')
                u_moo = first_row.get(cols['moo'], '-')
                u_phone = first_row.get(cols.get('phone', 'โทรศัพท์'), '-')

                num_licenses = len(group_data)

                with st.container(border=True):
                    # ส่วนหัวผู้ประกอบการ
                    col_h1, col_h2 = st.columns([5, 1])
                    
                    with col_h1:
                        st.markdown(f"""
                        <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 5px;">
                            <span style="font-size: 24px; color: #1E3A8A;">👤</span>
                            <h3 style="color: #1E3A8A; margin: 0;">{u_name}</h3>
                        </div>
                        <div style="color: #4B5563; font-size: 14px; margin-bottom: 10px; margin-left: 35px;">
                            💳 <b>นิติบุคคล/บัตรปชช:</b> {u_cid}<br>
                            📍 <b>ที่อยู่:</b> บ้านเลขที่ {u_addr} หมู่ที่ {u_moo} | ☎ <b>โทร:</b> {u_phone}
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col_h2:
                        with st.popover("🗑️ ลบข้อมูล", use_container_width=True):
                            st.warning(f"⚠️ ยืนยันการลบข้อมูลทั้งหมดของคุณ {u_name} หรือไม่? (รวม {num_licenses} รายการ)")
                            st.write("การลบนี้จะลบข้อมูลออกจาก Google Sheet โดยตรงและไม่สามารถกู้คืนได้")
                            if st.button("🔴 ยืนยันการลบทั้งหมด", key=f"confirm_del_{u_name}_{u_cid}", type="primary", use_container_width=True):
                                with st.spinner("กำลังลบข้อมูล..."):
                                    # คำนวณแถวที่ต้องลบ (df index + 2)
                                    rows_to_delete = [int(idx) + 2 for idx in group_data.index]
                                    if delete_gsheet_rows(rows_to_delete, sheet_name=target_sheet):
                                        st.success(f"ลบข้อมูลของ {u_name} เรียบร้อยแล้ว")
                                        time.sleep(1)
                                        st.cache_data.clear()
                                        st.rerun()
                    
                    if st.button(f"🖨️ พิมพ์ทุกใบอนุญาต ({num_licenses} ใบ)", key=f"print_all_{first_row.name}"):
                        st.session_state[f"do_print_all_{first_row.name}"] = True
                        
                    if st.session_state.get(f"do_print_all_{first_row.name}"):
                        with st.spinner("กำลังสร้างไฟล์ Word..."):
                            thai_months = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
                            
                            output_buffer = io.BytesIO()
                            with zipfile.ZipFile(output_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                for idx, row_item in group_data.iterrows():
                                    old_exp = row_item.get(cols['expire'])
                                    if pd.notnull(old_exp) and isinstance(old_exp, (pd.Timestamp, datetime)):
                                        try:
                                            issue_date = old_exp.replace(year=datetime.now().year)
                                        except ValueError:
                                            issue_date = old_exp.replace(year=datetime.now().year, day=28)
                                    else:
                                        issue_date = datetime.now()
                                        
                                    try:
                                        expire_date = issue_date.replace(year=issue_date.year + 1) - timedelta(days=1)
                                    except ValueError:
                                        expire_date = issue_date + timedelta(days=364)
                                    
                                    rcpt_date = datetime.now()
                                    
                                    context = {
                                        "p_license_book": "",
                                        "p_license_no": "",
                                        "p_license_year": str(datetime.now().year + 543),
                                        "p_name": str(row_item.get(cols['name'], '')),
                                        "p_nationality": "ไทย",
                                        "p_addr": str(row_item.get(cols['address'], '')),
                                        "p_moo": str(row_item.get(cols['moo'], '')),
                                        "p_tumbon": "ดอยงาม",
                                        "p_amphoe": "พาน",
                                        "p_province": "เชียงราย",
                                        "p_cid": str(row_item.get(cols['cid'], '')),
                                        "p_phone": str(row_item.get(cols.get('phone', 'โทรศัพท์'), '')),
                                        "p_shop": str(row_item.get(cols['shop'], '')),
                                        "p_type": str(row_item.get(cols['type'], '')),
                                        "p_fee": str(row_item.get(cols['fee'], '')),
                                        "p_fee_text": num_to_thai_baht(row_item.get(cols['fee'], '')) if row_item.get(cols['fee']) else "-",
                                        "p_rcpt_book": str(row_item.get(cols['rcpt_book'], '')),
                                        "p_rcpt_no": str(row_item.get(cols['rcpt_no'], '')),
                                        "p_rcpt_date": f"{rcpt_date.day} {thai_months[rcpt_date.month]} {rcpt_date.year + 543}",
                                        "issue_day": str(issue_date.day),
                                        "issue_month": thai_months[issue_date.month],
                                        "issue_year": str(issue_date.year + 543),
                                        "expire_day": str(expire_date.day),
                                        "expire_month": thai_months[expire_date.month],
                                        "expire_year": str(expire_date.year + 543),
                                        "p_43": str(row_item.get(cols.get('p_43', ''), '')),
                                        "p_44": str(row_item.get(cols.get('p_44', ''), '')),
                                    }
                                    docx_buf = create_docx_document(context)
                                    clean_name = str(row_item.get(cols['name'], 'ผู้ประกอบการ')).replace('/', '_').replace('\\', '_')
                                    clean_type = str(row_item.get(cols['type'], 'ใบอนุญาต')).replace('/', '_').replace('\\', '_')[:15]
                                    file_name = f"ใบอนุญาต_{clean_name}_{clean_type}_{idx}.docx"
                                    zip_file.writestr(file_name, docx_buf.getvalue())
                                    
                            output_buffer.seek(0)
                            
                        st.download_button(
                            label="📥 คลิกดาวน์โหลดไฟล์ Word ทั้งหมด (ZIP)",
                            data=output_buffer,
                            file_name=f"ใบอนุญาต_รวม_{u_name}.zip",
                            mime="application/zip",
                            type="primary",
                            key=f"dl_all_{first_row.name}"
                        )
                        
                    st.markdown("<hr style='margin: 10px 0; border-top: 1px solid #e5e7eb;'>", unsafe_allow_html=True)
                    
                    # ส่วนใบอนุญาตย่อย
                    for index, row in group_data.iterrows():
                        u_type = row.get(cols['type'], 'ไม่ระบุ')
                        u_shop = row.get(cols['shop'], '-')
                        u_fee = row.get(cols['fee'], '-')
                        u_expire = row.get(cols['expire'], '')
                        
                        rcpt_book = row.get(cols['rcpt_book'], '')
                        rcpt_no = row.get(cols['rcpt_no'], '')
                        u_rcpt = f"{rcpt_book}/{rcpt_no}" if rcpt_book and rcpt_no else "-"
                        
                        issue_date_str = "-"
                        expire_date_str = "-"
                        if isinstance(u_expire, pd.Timestamp) and not pd.isna(u_expire):
                            thai_year_expire = u_expire.year + 543
                            expire_date_str = f"{u_expire.day} {u_expire.strftime('%b')} {thai_year_expire}"
                            try:
                                issue_date = u_expire - timedelta(days=364)
                                thai_year_issue = issue_date.year + 543
                                issue_date_str = f"{issue_date.day} {issue_date.strftime('%b')} {thai_year_issue}"
                            except: pass
                        elif str(u_expire) != 'NaT':
                            expire_date_str = str(u_expire)

                        with st.container(border=True):
                            st.markdown(f"""
                            <div style="margin-bottom: 10px;">
                                <b style="color: #374151; font-size: 16px;">📄 {u_type}</b>
                            </div>
                            <div style="display: flex; gap: 10px; flex-wrap: wrap; margin-bottom: 10px;">
                                <span style="background-color: #fef3c7; color: #b45309; padding: 4px 10px; border-radius: 12px; font-size: 13px; border: 1px solid #fde68a;">💰 ค่าธรรมเนียม {u_fee} บาท</span>
                                <span style="background-color: #e0e7ff; color: #4338ca; padding: 4px 10px; border-radius: 12px; font-size: 13px; border: 1px solid #c7d2fe;">🗓️ ออกใบอนุญาต: {issue_date_str}</span>
                                <span style="background-color: #f3f4f6; color: #4b5563; padding: 4px 10px; border-radius: 12px; font-size: 13px; border: 1px solid #e5e7eb;">🗓️ หมดอายุ: {expire_date_str}</span>
                                <span style="background-color: #fce7f3; color: #be185d; padding: 4px 10px; border-radius: 12px; font-size: 13px; border: 1px solid #fbcfe8;">🏪 {u_shop}</span>
                            </div>
                            <div style="font-size: 13px; color: #6b7280; margin-bottom: 10px;">ใบเสร็จ: {u_rcpt}</div>
                            """, unsafe_allow_html=True)
                            
                            c1, c2, c3, c4 = st.columns([1.5, 1.5, 1.5, 5])
                            if c1.button("🔄 ต่ออายุ", key=f"ed_{index}", use_container_width=True): st.session_state[f"m_{index}"] = "edit"
                            if c2.button("📜 พิมพ์ใบอนุญาต", key=f"pr_{index}", use_container_width=True): st.session_state[f"m_{index}"] = "print"
                            uploaded_files = c3.file_uploader("📁 สแกน", key=f"up_{index}", label_visibility="collapsed", accept_multiple_files=True)
                            if uploaded_files:
                                if st.button("📤 อัปโหลด", key=f"btn_up_{index}", use_container_width=True):
                                    # แถวใน Sheet คือ index + 2
                                    sheet_row = int(index) + 2
                                    success = True
                                    for f in uploaded_files:
                                        if not upload_to_gdrive(f, person_name, sheet_row, target_sheet):
                                            success = False
                                    if success:
                                        st.success("✅ อัปโหลดไฟล์ทั้งหมดสำเร็จ!")
                                        time.sleep(1)
                                        st.cache_data.clear()
                                        st.rerun()
                            
                            if st.session_state.get(f"m_{index}") in ["print", "edit"]:
                                if st.button("✖️ ปิดแถบจัดการข้อมูล", key=f"close_{index}", use_container_width=True):
                                    st.session_state[f"m_{index}"] = None
                                    st.rerun()
                                    
                            if st.session_state.get(f"m_{index}") == "print":
                                tab1, tab2 = st.tabs(["📄 ใบอนุญาต (อภ.๒)", "📝 คำขอ"])
                                
                                with tab1:
                                    st.subheader("จัดการข้อมูลและแก้ไขก่อนพิมพ์ (อภ.๒)")
                                    
                                    # ดึงค่าจากชีตมาตรงๆ ไม่ต้องคำนวณใหม่ (ตามคำขอผู้ใช้)
                                    old_exp = row.get(cols['expire'])
                                    old_issue = row.get(cols.get('rcpt_date', 'ลงวันที่'))
                                    
                                    # แปลงเป็น datetime object เพื่อใช้กับ date_input
                                    try:
                                        if pd.notna(old_issue) and isinstance(old_issue, (pd.Timestamp, datetime)):
                                            issue_default = old_issue.date()
                                        else:
                                            # ลองแปลงจาก String ถ้าไม่ใช่ Timestamp
                                            issue_default = datetime.strptime(str(old_issue).split(' ')[0], '%d/%m/%Y').date()
                                    except:
                                        issue_default = datetime.now().date()
                                        
                                    try:
                                        if pd.notna(old_exp) and isinstance(old_exp, (pd.Timestamp, datetime)):
                                            expire_default = old_exp.date()
                                        else:
                                            # ลองแปลงจาก String ถ้าไม่ใช่ Timestamp
                                            expire_default = datetime.strptime(str(old_exp).split(' ')[0], '%d/%m/%Y').date()
                                    except:
                                        # ถ้าคำนวณไม่ได้ ให้เป็น +1 ปี จากวันออก
                                        expire_default = (issue_default + timedelta(days=365))
                                        try:
                                            expire_default = expire_default.replace(day=issue_default.day) - timedelta(days=1)
                                        except: pass
                                    
                                    with st.container(border=True):
                                        col_f1, col_f2 = st.columns(2)
                                        
                                        p_name = col_f1.text_input("ชื่อผู้รับใบอนุญาต", value=row.get(cols['name'], ''), key=f"p_name_{index}")
                                        p_nationality = col_f2.text_input("สัญชาติ", value="ไทย", key=f"p_nat_{index}")
                                        
                                        p_cid = col_f1.text_input("เลขประจำตัวประชาชน/นิติบุคคล", value=row.get(cols['cid'], ''), key=f"p_cid_{index}")
                                        p_phone = col_f2.text_input("โทรศัพท์", value=row.get(cols.get('phone', 'โทรศัพท์'), ''), key=f"p_phone_{index}")
                                        
                                        p_addr = col_f1.text_input("ที่อยู่/บ้านเลขที่", value=row.get(cols['address'], ''), key=f"p_addr_{index}")
                                        p_moo = col_f2.text_input("หมู่ที่", value=row.get(cols['moo'], ''), key=f"p_moo_{index}")
                                        
                                        c_addr_o1, c_addr_o2, c_addr_o3 = st.columns(3)
                                        p_tumbon = c_addr_o1.text_input("ตำบล", value="ดอยงาม", key=f"p_t_{index}")
                                        p_amphoe = c_addr_o2.text_input("อำเภอ", value="พาน", key=f"p_a_{index}")
                                        p_province = c_addr_o3.text_input("จังหวัด", value="เชียงราย", key=f"p_p_{index}")
                                        
                                        st.markdown("---")
                                        st.markdown("**ข้อมูลสถานประกอบการ**")
                                        
                                        col_s1, col_s2 = st.columns(2)
                                        p_shop = col_s1.text_input("ชื่อสถานประกอบการ", value=row.get(cols['shop'], ''), key=f"p_shop_{index}")
                                        p_type = col_s2.text_input("ประเภทการทำกิจการ", value=row.get(cols['type'], ''), key=f"p_type_{index}")
                                        
                                        p_shop_addr = col_s1.text_input("ที่อยู่สถานประกอบการ", value=row.get(cols['address'], ''), key=f"p_s_addr_{index}")
                                        p_shop_moo = col_s2.text_input("หมู่ที่ (สถานประกอบการ)", value=row.get(cols['moo'], ''), key=f"p_s_moo_{index}")
                                        
                                        c_addr_s1, c_addr_s2, c_addr_s3 = st.columns(3)
                                        p_shop_tumbon = c_addr_s1.text_input("ตำบล (ร้าน)", value="ดอยงาม", key=f"p_s_t_{index}")
                                        p_shop_amphoe = c_addr_s2.text_input("อำเภอ (ร้าน)", value="พาน", key=f"p_s_a_{index}")
                                        p_shop_province = c_addr_s3.text_input("จังหวัด (ร้าน)", value="เชียงราย", key=f"p_s_p_{index}")
                                        
                                        p_phone = col_s2.text_input("โทรศัพท์ (ร้าน)", value=row.get(cols.get('phone', 'โทรศัพท์'), ''), key=f"p_s_phone_{index}")
                                        
                                        st.markdown("---")
                                        st.markdown("**ข้อมูลใบอนุญาต และใบเสร็จรับเงิน**")
                                        
                                        c_l1, c_l2, c_l3 = st.columns(3)
                                        p_license_book = c_l1.text_input("ใบอนุญาตเล่มที่", value="", key=f"p_l_book_{index}")
                                        p_license_no = c_l2.text_input("ใบอนุญาตเลขที่", value="", key=f"p_l_no_{index}")
                                        p_license_year = c_l3.text_input("ปี (พ.ศ.)", value=str(datetime.now().year + 543), key=f"p_l_year_{index}")
                                        
                                        c_r1, c_r2, c_r3, c_r4 = st.columns(4)
                                        fee_val = str(row.get(cols['fee'], ''))
                                        default_fee_txt = num_to_thai_baht(fee_val)
                                        
                                        p_fee = c_r1.text_input(
                                            "ค่าธรรมเนียม (ตัวเลข)", 
                                            value=fee_val, 
                                            key=f"p_fee_{index}",
                                            on_change=update_fee_text,
                                            args=(index,)
                                        )
                                        p_fee_text = c_r2.text_input(
                                            "ค่าธรรมเนียม (ตัวอักษร)", 
                                            value=default_fee_txt, 
                                            key=f"p_fee_txt_{index}"
                                        )
                                        p_rcpt_book = c_r3.text_input("ใบเสร็จเล่มที่", value=str(row.get(cols['rcpt_book'], '')), key=f"p_book_{index}")
                                        p_rcpt_no = c_r4.text_input("ใบเสร็จเลขที่", value=str(row.get(cols['rcpt_no'], '')), key=f"p_no_{index}")
                                        
                                        c_d1, c_d2, c_d3 = st.columns(3)
                                        # (New condition fields 4.3 and 4.4 added below)
                                        st.markdown("---")
                                        st.markdown("**เงื่อนไขเพิ่มเติม (ข้อ 4)**")
                                        c_41, c_42 = st.columns(2)
                                        p_43 = c_41.text_input("เงื่อนไขเพิ่มเติม 4.3", value=row.get(cols.get('p_43', ''), ''), key=f"p_43_{index}")
                                        p_44 = c_42.text_input("เงื่อนไขเพิ่มเติม 4.4", value=row.get(cols.get('p_44', ''), ''), key=f"p_44_{index}")
                                        st.markdown("---")
                                        p_rcpt_date = c_d1.date_input("ลงวันที่ (ใบเสร็จ)", value=issue_default, key=f"p_rcpt_date_{index}")
                                        p_issue = c_d2.date_input("วันที่ออกใบอนุญาต", value=issue_default, key=f"p_issue_{index}")
                                        p_expire = c_d3.date_input("วันหมดอายุ", value=expire_default, key=f"p_expire_{index}")

                                    if st.button("💾 บันทึกและสร้างไฟล์ PDF (อภ.๒)", key=f"preview_btn_{index}", type="primary"):
                                        thai_months = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
                                        
                                        # Prepare data for Google Sheet Update
                                        update_data = {
                                            cols['name']: p_name,
                                            cols['cid']: p_cid,
                                            cols['shop']: p_shop,
                                            cols['type']: p_type,
                                            cols['address']: p_addr,
                                            cols['moo']: p_moo,
                                            cols.get('phone', 'โทรศัพท์'): p_phone,
                                            cols['fee']: p_fee,
                                            cols['rcpt_book']: p_rcpt_book,
                                            cols['rcpt_no']: p_rcpt_no,
                                            cols['rcpt_date']: p_rcpt_date.strftime('%d/%m/%Y'),
                                            cols['expire']: p_expire.strftime('%d/%m/%Y')
                                        }
                                        if cols.get('p_43') and cols['p_43'] in df.columns:
                                            update_data[cols['p_43']] = p_43
                                        if cols.get('p_44') and cols['p_44'] in df.columns:
                                            update_data[cols['p_44']] = p_44
                                        
                                        # Save to Google Sheet
                                        with st.spinner("กำลังบันทึกข้อมูลลง Google Sheet..."):
                                            if update_gsheet(index + 2, update_data, sheet_name=target_sheet):
                                                st.toast("✅ บันทึกข้อมูลลงชีตเรียบร้อยแล้ว", icon="💾")
                                            else:
                                                st.error("⚠️ บันทึกข้อมูลลงชีตไม่สำเร็จ แต่จะดำเนินการสร้าง PDF ต่อ")

                                        context = {
                                            "p_license_book": p_license_book,
                                            "p_license_no": p_license_no,
                                            "p_license_year": p_license_year,
                                            "p_name": p_name,
                                            "p_nationality": p_nationality,
                                            "p_addr": p_addr,
                                            "p_moo": p_moo,
                                            "p_tumbon": p_tumbon,
                                            "p_amphoe": p_amphoe,
                                            "p_province": p_province,
                                            "p_cid": p_cid,
                                            "p_phone": p_phone,
                                            "p_shop": p_shop,
                                            "p_type": p_type,
                                            "p_shop_addr": p_shop_addr,
                                            "p_shop_moo": p_shop_moo,
                                            "p_shop_tumbon": p_shop_tumbon,
                                            "p_shop_amphoe": p_shop_amphoe,
                                            "p_shop_province": p_shop_province,
                                            "p_shop_phone": p_phone,
                                            "p_fee": p_fee,
                                            "p_fee_text": p_fee_text,
                                            "p_rcpt_book": p_rcpt_book,
                                            "p_rcpt_no": p_rcpt_no,
                                            "p_rcpt_date": f"{p_rcpt_date.day} {thai_months[p_rcpt_date.month]} {p_rcpt_date.year + 543}",
                                            "issue_day": str(p_issue.day),
                                            "issue_month": thai_months[p_issue.month],
                                            "issue_year": str(p_issue.year + 543),
                                            "expire_day": str(p_expire.day),
                                            "expire_month": thai_months[p_expire.month],
                                            "expire_year": str(p_expire.year + 543),
                                            "p_43": p_43,
                                            "p_44": p_44,
                                        }
                                        try:
                                            docx_buffer = create_docx_document(context)
                                            
                                            has_special = bool(str(p_43).strip() or str(p_44).strip())
                                            
                                            if not has_special:
                                                pdf_buffer = create_pdf_overlay(context)
                                                col1, col2 = st.columns(2)
                                                with col1:
                                                    st.download_button(
                                                        label="📥 ดาวน์โหลดไฟล์ใบอนุญาต (Word)",
                                                        data=docx_buffer,
                                                        file_name=f"ใบอนุญาต_{p_name}.docx",
                                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                        type="secondary",
                                                        key=f"dl_lic_btn_word_{index}",
                                                        use_container_width=True
                                                    )
                                                with col2:
                                                    if pdf_buffer:
                                                        st.download_button(
                                                            label="📥 ดาวน์โหลดไฟล์ใบอนุญาต (PDF)",
                                                            data=pdf_buffer,
                                                            file_name=f"ใบอนุญาต_{p_name}.pdf",
                                                            mime="application/pdf",
                                                            type="primary",
                                                            key=f"dl_lic_btn_pdf_{index}",
                                                            use_container_width=True
                                                        )
                                                    else:
                                                        st.error("ไม่สามารถสร้างไฟล์ PDF ใบอนุญาตได้เนื่องจากไม่พบไฟล์เทมเพลต")
                                                st.success("สร้างไฟล์ใบอนุญาตสำเร็จแล้ว! สามารถเลือกดาวน์โหลดรูปแบบ Word หรือ PDF ได้")
                                            else:
                                                st.download_button(
                                                    label="📥 ดาวน์โหลดไฟล์ใบอนุญาต (Word)",
                                                    data=docx_buffer,
                                                    file_name=f"ใบอนุญาต_{p_name}.docx",
                                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                    type="secondary",
                                                    key=f"dl_lic_btn_word_only_{index}",
                                                    use_container_width=True
                                                )
                                                st.info("💡 เนื่องจากมีเงื่อนไขเพิ่มเติมระบบจึงสร้างเฉพาะไฟล์ Word เพื่อให้คุณปรับแต่งข้อความเพิ่มเติมได้สะดวกครับ")
                                                st.success("สร้างไฟล์ Word ใบอนุญาตพร้อมแล้ว กรุณากดปุ่มดาวน์โหลดด้านบน ↑")
                                        except Exception as e:
                                            st.error(f"เกิดข้อผิดพลาดในการสร้างไฟล์ใบอนุญาต: {e}")
                                            
                                    with tab2:
                                        st.subheader("จัดการข้อมูลและแก้ไขก่อนพิมพ์ (แบบคำขอ)")
                                        
                                        with st.container(border=True):
                                            col_a1, col_a2 = st.columns(2)
                                            
                                            a_type = col_a1.text_input("ประเภทการทำกิจการ", value=row.get(cols['type'], ''), key=f"a_type_{index}")
                                            a_name = col_a2.text_input("ข้าพเจ้า (ชื่อผู้ขอ)", value=row.get(cols['name'], ''), key=f"a_name_{index}")
                                            
                                            a_age = col_a1.text_input("อายุ (ปี)", value="", key=f"a_age_{index}")
                                            a_nat = col_a2.text_input("สัญชาติ", value="ไทย", key=f"a_nat_{index}")
                                            
                                            a_agent = col_a1.text_input("โดย (ผู้มีอำนาจลงนามแทนนิติบุคคล)", value="", key=f"a_agent_{index}")
                                            a_evidence = col_a2.text_input("ปรากฏตาม", value="", placeholder="ระบุเอกสารอ้างอิง เช่น บัตรประจำตัวประชาชน", key=f"a_evid_{index}")
                                            
                                            a_phone = col_a1.text_input("หมายเลขโทรศัพท์", value=row.get(cols.get('phone', 'โทรศัพท์'), ''), key=f"a_phone_{index}")
                                            
                                            st.markdown("**ที่อยู่ผู้ขออนุญาต**")
                                            c_addr1, c_addr2, c_addr3, c_addr4 = st.columns(4)
                                            a_addr = c_addr1.text_input("ที่อยู่เลขที่", value=row.get(cols['address'], ''), key=f"a_addr_{index}")
                                            a_moo = c_addr2.text_input("หมู่ที่", value=row.get(cols['moo'], ''), key=f"a_moo_{index}")
                                            a_soi = c_addr3.text_input("ตรอก/ซอย", value="", key=f"a_soi_{index}")
                                            a_road = c_addr4.text_input("ถนน", value="", key=f"a_road_{index}")
                                            
                                            c_addr5, c_addr6, c_addr7 = st.columns(3)
                                            a_subdist = c_addr5.text_input("แขวง/ตำบล", value="ดอยงาม", key=f"a_subdist_{index}")
                                            a_dist = c_addr6.text_input("เขต/อำเภอ", value="พาน", key=f"a_dist_{index}")
                                            a_prov = c_addr7.text_input("จังหวัด", value="เชียงราย", key=f"a_prov_{index}")
                                            
                                            st.markdown("**ข้อมูลการเขียนคำขอ**")
                                            c_req1, c_req2 = st.columns(2)
                                            a_location = c_req1.text_input("เขียนที่", value="อบต.ดอยงาม", key=f"a_loc_{index}")
                                            a_date = c_req2.date_input("วันที่ยื่นคำขอ", value=datetime.now(), key=f"a_date_{index}")
                                            
                                            st.markdown("**ข้อ ๒. เอกสารหลักฐานที่แนบมาด้วย**")
                                            chk_1 = st.checkbox("สำเนาบัตรประจำตัว", value=False, key=f"chk_1_{index}")
                                            chk_2 = st.checkbox("สำเนาใบอนุญาตตามกฎหมายที่เกี่ยวข้อง", value=False, key=f"chk_2_{index}")
                                            chk_3 = st.checkbox("หนังสือให้ความเห็นชอบการประเมินผลกระทบต่อสิ่งแวดล้อม", value=False, key=f"chk_3_{index}")
                                            chk_4 = st.checkbox("ใบมอบอำนาจ", value=False, key=f"chk_4_{index}")
                                            chk_5 = st.checkbox("สำเนาหนังสือรับรองการจดทะเบียนเป็นนิติบุคคล", value=False, key=f"chk_5_{index}")
                                            chk_6 = st.checkbox("หลักฐานที่แสดงการเป็นผู้มีอำนาจลงนามแทนนิติบุคคล", value=False, key=f"chk_6_{index}")
                                            chk_7 = st.checkbox("เอกสารและหลักฐานอื่นๆ", value=False, key=f"chk_7_{index}")
                                            
                                        if st.button("💾 บันทึกและสร้างไฟล์ PDF (คำขอ)", key=f"app_btn_{index}", type="primary"):
                                            thai_months = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
                                            
                                            # Prepare data for Google Sheet Update
                                            update_data = {
                                                cols['name']: a_name,
                                                cols['type']: a_type,
                                                cols['address']: a_addr,
                                                cols['moo']: a_moo,
                                                cols.get('phone', 'โทรศัพท์'): a_phone
                                            }
                                            
                                            # Save to Google Sheet
                                            with st.spinner(f"กำลังบันทึกข้อมูลลงชีต {target_sheet}..."):
                                                if update_gsheet(index + 2, update_data, sheet_name=target_sheet):
                                                    st.toast("✅ บันทึกข้อมูลลงชีตเรียบร้อยแล้ว", icon="💾")
                                                else:
                                                    st.error("⚠️ บันทึกข้อมูลลงชีตไม่สำเร็จ แต่จะดำเนินการสร้าง PDF ต่อ")

                                            app_context = {
                                                "p_type": a_type,
                                                "app_location": a_location,
                                                "app_day": str(a_date.day),
                                                "app_month": thai_months[a_date.month],
                                                "app_year": str(a_date.year + 543),
                                                "p_name": a_name,
                                                "p_age": a_age,
                                                "p_nationality": a_nat,
                                                "p_agent": a_agent,
                                                "p_evidence": a_evidence,
                                                "p_addr": a_addr,
                                                "p_moo": a_moo,
                                                "p_soi": a_soi,
                                                "p_road": a_road,
                                                "p_subdistrict": a_subdist,
                                                "p_district": a_dist,
                                                "p_province": a_prov,
                                                "p_phone": a_phone,
                                                "chk_1": chk_1,
                                                "chk_2": chk_2,
                                                "chk_3": chk_3,
                                                "chk_4": chk_4,
                                                "chk_5": chk_5,
                                                "chk_6": chk_6,
                                                "chk_7": chk_7
                                            }
                                            
                                            try:
                                                app_docx_buffer = create_app_docx_document(app_context)
                                                app_pdf_buffer = create_app_pdf_overlay(app_context)
                                                
                                                col1, col2 = st.columns(2)
                                                with col1:
                                                    st.download_button(
                                                        label="📥 ดาวน์โหลดไฟล์คำขอ (Word)",
                                                        data=app_docx_buffer,
                                                        file_name=f"แบบคำขอ_{a_name}.docx",
                                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                        type="secondary",
                                                        key=f"dl_app_btn_word_{index}",
                                                        use_container_width=True
                                                    )
                                                with col2:
                                                    if app_pdf_buffer:
                                                        st.download_button(
                                                            label="📥 ดาวน์โหลดไฟล์คำขอ (PDF)",
                                                            data=app_pdf_buffer,
                                                            file_name=f"แบบคำขอ_{a_name}.pdf",
                                                            mime="application/pdf",
                                                            type="primary",
                                                            key=f"dl_app_btn_pdf_{index}",
                                                            use_container_width=True
                                                        )
                                                    else:
                                                        st.error("ไม่สามารถสร้างไฟล์ PDF คำขอได้เนื่องจากไม่พบไฟล์เทมเพลต")
                                                st.success("สร้างไฟล์คำขอสำเร็จแล้ว! คุณสามารถเลือกดาวน์โหลดได้ทั้งแบบ Word (สำหรับนำไปแก้ไข) และ PDF (สำหรับสั่งพิมพ์)")
                                            except Exception as e:
                                                st.error(f"เกิดข้อผิดพลาดในการสร้างไฟล์คำขอ: {e}")
                            
                            elif st.session_state.get(f"m_{index}") == "edit":
                                st.subheader("📝 แก้ไขข้อมูล / ต่ออายุ")
                                with st.form(f"edit_form_{index}"):
                                    c_e1, c_e2 = st.columns(2)
                                    
                                    edit_data = {}
                                    # กรองเอาเฉพาะคอลัมน์ที่มีชื่อ และไม่เป็น Unnamed หรือคอลัมน์ไฟล์
                                    visible_cols = [c for c in df.columns if c and not c.startswith('Unnamed') and c not in ['ไฟล์', 'ไฟล์แนบ']]
                                    
                                    for i, col_name in enumerate(visible_cols):
                                        col_ui = c_e1 if i % 2 == 0 else c_e2
                                        val = row.get(col_name, "")
                                        if pd.isna(val) or str(val).lower() == "nan": val = ""
                                        
                                        if 'วัน' in col_name:
                                            # Special logic for auto-calculating expiration date for renewal
                                            if col_name == cols.get('expire', 'วันหมดอายุใบอนุญาต') or col_name == 'วันหมดอายุใบอนุญาต':
                                                try:
                                                    if isinstance(val, (pd.Timestamp, datetime)):
                                                        orig_date = val.date()
                                                    elif val and '-' in str(val):
                                                        orig_date = datetime.strptime(str(val).split(' ')[0], '%Y-%m-%d').date()
                                                    elif val:
                                                        orig_date = datetime.strptime(str(val).split(' ')[0], '%d/%m/%Y').date()
                                                    else:
                                                        orig_date = datetime.now().date()
                                                        
                                                    # Issue date = original date but current year
                                                    try:
                                                        issue_d = orig_date.replace(year=datetime.now().year)
                                                    except ValueError:
                                                        issue_d = orig_date.replace(year=datetime.now().year, day=28)
                                                        
                                                    # Expire date = issue date + 1 year - 1 day
                                                    try:
                                                        default_date = issue_d.replace(year=issue_d.year + 1) - timedelta(days=1)
                                                    except ValueError:
                                                        default_date = issue_d.replace(year=issue_d.year + 1, day=28) - timedelta(days=1)
                                                except:
                                                    default_date = datetime.now().date()
                                            elif col_name == cols.get('rcpt_date', 'ลงวันที่') or col_name == 'ลงวันที่':
                                                default_date = datetime.now().date()
                                            else:
                                                try:
                                                    if isinstance(val, pd.Timestamp):
                                                        default_date = val.date()
                                                    elif isinstance(val, datetime):
                                                        default_date = val.date()
                                                    elif val:
                                                        if '-' in str(val):
                                                            default_date = datetime.strptime(str(val).split(' ')[0], '%Y-%m-%d').date()
                                                        else:
                                                            default_date = datetime.strptime(str(val).split(' ')[0], '%d/%m/%Y').date()
                                                    else:
                                                        default_date = datetime.now().date()
                                                except:
                                                    default_date = datetime.now().date()
                                            
                                            edit_data[col_name] = col_ui.date_input(col_name, value=default_date).strftime('%d/%m/%Y')
                                        else:
                                            edit_data[col_name] = col_ui.text_input(col_name, value=str(val))
                                            
                                    if st.form_submit_button("💾 บันทึกการแก้ไขลงชีต"):
                                        with st.spinner(f"กำลังอัปเดตข้อมูลในชีต {target_sheet}..."):
                                            if update_gsheet(index + 2, edit_data, sheet_name=target_sheet):
                                                st.success("✅ บันทึกข้อมูลสำเร็จ!")
                                                time.sleep(1)
                                                st.cache_data.clear()
                                                st.rerun()
                                            else:
                                                st.error("❌ บันทึกไม่สำเร็จ กรุณาลองใหม่อีกครั้ง")
        else: st.warning("❌ ไม่พบข้อมูล")

